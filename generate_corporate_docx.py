import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_header_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1A56DB')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_corporate_guide():
    doc = Document()
    
    # Define Corporate Colors
    brand_blue = RGBColor(26, 86, 219)     # #1A56DB
    dark_slate = RGBColor(15, 23, 42)      # #0F172A
    muted_text = RGBColor(71, 85, 105)     # #475569
    
    # Modify default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Title
    title = doc.add_heading('RETRIX ENTERPRISE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = dark_slate
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.bold = True
    
    subtitle = doc.add_paragraph('OFFICIAL TRAINING GUIDE')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.runs[0]
    sub_run.font.color.rgb = brand_blue
    sub_run.font.size = Pt(18)
    sub_run.font.bold = True
    
    doc.add_paragraph('\n')
    
    # Metadata block
    meta = doc.add_paragraph()
    meta.add_run('Prepared By: ').bold = True
    meta.add_run('Simwinga Felix\n')
    meta.add_run('Audience: ').bold = True
    meta.add_run('Retrix Enterprise Administration & Staff\n')
    meta.add_run('Support/Query: ').bold = True
    meta.add_run('simwinga8788@gmail.com\n')
    meta.add_run('Phone: ').bold = True
    meta.add_run('0972996902')
    
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # Helper to add corporate H1
    def add_h1(text):
        h = doc.add_heading(text, level=1)
        r = h.runs[0]
        r.font.color.rgb = brand_blue
        r.font.size = Pt(16)
        r.font.bold = True
        add_header_border(h)
        return h

    # Helper to add corporate H2
    def add_h2(text):
        h = doc.add_heading(text, level=2)
        r = h.runs[0]
        r.font.color.rgb = dark_slate
        r.font.size = Pt(13)
        r.font.bold = True
        return h
    
    # 1. Introduction
    add_h1('1. Introduction to the Platform')
    doc.add_paragraph('Welcome to the Retrix Car Rental Management System. This platform has been engineered to streamline every aspect of our car rental business—from the moment a customer views our fleet online, to booking management, payment collection at the counter, and vehicle return.')
    doc.add_paragraph('This guide provides a comprehensive walkthrough of the system to ensure you and your team can operate the platform efficiently and professionally.')
    
    # 2. Customer Experience
    add_h1('2. The Customer Experience (Public Portal)')
    doc.add_paragraph('Before managing the backend, it is crucial to understand what our customers see and how they interact with Retrix.')
    
    add_h2('2.1 Browsing the Fleet')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Dynamic Fleet Viewing: ').bold = True
    p.add_run('Customers navigate to the Fleet page to browse available vehicles. The system automatically tags vehicles with key features (e.g., Automatic, 4 Seater, AC).')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Search & Filters: ').bold = True
    p.add_run('Customers can actively filter vehicles by transmission type, price range, and category to find exactly what they need.')
    
    add_h2('2.2 The Booking Process (Pay at Counter)')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Vehicle Selection: ').bold = True
    p.add_run('The customer selects a vehicle to view detailed specifications, pricing (Daily/Weekly rates), and high-resolution images.')
    
    p = doc.add_paragraph(style='List Number')
    p.add_run('Date Selection: ').bold = True
    p.add_run('The built-in calendar ensures customers cannot book a car that is already reserved or in maintenance for the selected dates.')
    
    p = doc.add_paragraph(style='List Number')
    p.add_run('Quotation Generation: ').bold = True
    p.add_run('Upon submitting a booking request, the system instantly generates a professional Quotation detailing the total cost (in ZMW and USD), rental duration, and the specific vehicle.')
    
    p = doc.add_paragraph(style='List Number')
    p.add_run('Counter Payment & Pickup: ').bold = True
    p.add_run('Customers reserve the vehicle online and finalize the payment physically at the counter when they arrive to pick up the car.')
    
    # 3. Admin Dashboard
    add_h1('3. The Administrative Dashboard')
    doc.add_paragraph('The Admin Portal is the central nervous system of Retrix Enterprise. Only users with designated Admin privileges can access this area.')
    
    add_h2('3.1 Analytics Dashboard')
    doc.add_paragraph('Upon logging in, you are greeted by the Analytics Dashboard. This provides a real-time, birds-eye view of business health:')
    doc.add_paragraph('Total Earnings: Gross revenue calculations based on completed bookings.', style='List Bullet')
    doc.add_paragraph('Fleet Utilization: A live percentage showing how much of the fleet is currently rented out versus sitting idle.', style='List Bullet')
    doc.add_paragraph('Interactive Charts: Visual breakdowns of revenue trends over the last 6 months and current fleet status distributions.', style='List Bullet')
    
    add_h2('3.2 Exporting Professional Reports')
    doc.add_paragraph('Action: Click the gold "Export to Excel" button in the top right corner.', style='List Bullet')
    doc.add_paragraph('Dynamic Timeframes: Choose between Daily, Weekly, Monthly, or Yearly reports.', style='List Bullet')
    doc.add_paragraph('Output: The system instantly generates a boardroom-ready Excel spreadsheet. The data is pre-calculated for your chosen timeframe, detailing Top Performing Vehicles, Total Revenue, and Recent Bookings.', style='List Bullet')
    
    # 4. Managing Bookings
    add_h1('4. Managing Bookings & Payments')
    add_h2('4.1 The Booking Lifecycle')
    doc.add_paragraph('Every booking moves through a strict status pipeline:')
    doc.add_paragraph('Pending: The customer has requested the car online, but has not yet arrived at the counter to pay.', style='List Number')
    doc.add_paragraph('Confirmed: The customer has arrived, paid at the counter, and the car is officially assigned to them.', style='List Number')
    doc.add_paragraph('Active: The customer has picked up the vehicle and driven off the lot.', style='List Number')
    doc.add_paragraph('Completed: The customer has returned the vehicle safely.', style='List Number')
    doc.add_paragraph('Cancelled: The booking was voided due to a no-show or customer cancellation.', style='List Number')
    
    add_h2('4.2 Processing Counter Payments')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Action: ').bold = True
    p.add_run('When a customer arrives at the counter, locate their Pending booking in the Bookings or Payments tab.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Action: ').bold = True
    p.add_run('Collect the physical or card payment. Once funds are secured, update the system to mark the payment as Completed. This immediately authorizes the release of the vehicle.')
    
    # 5. Fleet Management
    add_h1('5. Fleet Management & Damage Tracking')
    add_h2('5.1 Fleet Directory')
    doc.add_paragraph('Navigate to the Fleet tab to see every vehicle owned by Retrix.', style='List Bullet')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Adding a Vehicle: ').bold = True
    p.add_run('Click "Add Vehicle". You must provide the Make, Model, Transmission, Seats, and the Daily Rental Price. You can upload an image URL to represent the car.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Status Controls: ').bold = True
    p.add_run('You can manually set a car to Available, In Maintenance, or Unavailable. Note: A car automatically becomes Rented when a booking becomes Active.')
    
    add_h2('5.2 Damage & Incident Reports')
    doc.add_paragraph('Accidents and wear-and-tear happen. The Damages tab ensures no financial loss slips through the cracks.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Logging an Incident: ').bold = True
    p.add_run('If a car is returned damaged, log it immediately. Select the vehicle, input the date of the incident, and describe the damage (e.g., "Scratched left bumper").')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Financial Tracking: ').bold = True
    p.add_run('Input the Estimated Cost and, once fixed, the Actual Cost. This data is critical as it feeds directly into the Excel Reports to calculate net profitability per vehicle.')
    
    # 6. Customer Management
    add_h1('6. Customer Management')
    doc.add_paragraph('The Customers tab provides a complete CRM (Customer Relationship Management) view.')
    doc.add_paragraph('User Overview: See every registered user, their contact information, and how many bookings they have made.', style='List Bullet')
    doc.add_paragraph('Access Control (Admin): You can elevate trusted staff members by clicking the Admin shield icon next to their name. This grants them full access to the backend portal.', style='List Bullet')
    doc.add_paragraph('Disciplinary Action (Suspend): If a customer violates rental terms (e.g., late returns, reckless driving), you can click the red Suspend button. Suspended users are immediately blocked from making any future bookings on the platform.', style='List Bullet')
    
    # 7. System Settings & Locations
    add_h1('7. System Settings & Locations')
    add_h2('7.1 Managing Locations')
    doc.add_paragraph('Navigate to the Locations tab to manage pickup and drop-off points.', style='List Bullet')
    doc.add_paragraph('Having accurate locations ensures customers know exactly where to retrieve their vehicles.', style='List Bullet')
    
    add_h2('7.2 Core Settings')
    doc.add_paragraph('The Settings page allows you to control global variables such as the base currency exchange rates. Ensuring this is accurate guarantees that USD quotations are generated correctly based on the current ZMW rate.', style='List Bullet')
    
    # 8. Support
    add_h1('8. Support & Escalation')
    doc.add_paragraph('This platform is built for stability and ease of use. However, if you encounter technical anomalies, require new feature implementations, or need system maintenance, please escalate the issue using the contact details below.')
    
    contact = doc.add_paragraph()
    contact.add_run('Technical Lead: ').bold = True
    contact.add_run('Simwinga Felix\n')
    contact.add_run('Email: ').bold = True
    contact.add_run('simwinga8788@gmail.com\n')
    contact.add_run('Direct Line: ').bold = True
    contact.add_run('0972996902')
    
    save_path = os.path.join(os.environ['USERPROFILE'], 'Desktop', 'Retrix_Corporate_Guide.docx')
    doc.save(save_path)
    print(f'Corporate document saved to {save_path}')

if __name__ == "__main__":
    create_corporate_guide()
